"""
DoneGuard Lite
--------------
A single-page Streamlit app that helps Agile Release Trains (SAFe) enforce
Built-In Quality by automatically generating a Feature-level Definition of
Done (DoD) checklist, required sign-offs, and draft release notes — powered
by the Groq API (OpenAI-compatible chat completions, free tier, no card
required).

Run with:
    streamlit run app.py
"""

import io
import re

import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor


# ---------------------------------------------------------------------------
# Page configuration
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="DoneGuard Lite",
    page_icon="🛡️",
    layout="centered",
)

# ---------------------------------------------------------------------------
# Visual design system — a focused palette/type choice instead of Streamlit's
# default look. Deep navy ink + paper background + emerald "verified" accent,
# a serif display face for the title, IBM Plex Sans for body/UI text.
# ---------------------------------------------------------------------------
INK = "#16233D"
GUARD_GREEN = "#1F7A5C"
FLAG_AMBER = "#C97A1A"
PAPER_BG = "#F7F6F2"
BORDER = "#DEDAD0"

st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Fraunces:opsz,wght@9..144,600;9..144,700&family=IBM+Plex+Sans:wght@400;500;600&display=swap');

html, body, [class*="css"] {{
    font-family: 'IBM Plex Sans', sans-serif;
}}
.stApp {{ background-color: {PAPER_BG}; }}
section[data-testid="stSidebar"] {{
    background-color: #EFEDE5;
    border-right: 1px solid {BORDER};
}}
div[data-testid="stButton"] button {{
    font-family: 'IBM Plex Sans', sans-serif;
    border-radius: 4px;
    border: 1px solid {INK};
    color: {INK};
    background-color: transparent;
    font-weight: 500;
}}
div[data-testid="stButton"] button:hover {{
    border-color: {GUARD_GREEN};
    color: {GUARD_GREEN};
}}
div[data-testid="stButton"] button[kind="primary"] {{
    background-color: {GUARD_GREEN};
    border-color: {GUARD_GREEN};
    color: {PAPER_BG};
}}
div[data-testid="stButton"] button[kind="primary"]:hover {{
    background-color: #185F49;
    border-color: #185F49;
}}
textarea {{
    border-radius: 4px !important;
    border: 1px solid {BORDER} !important;
    font-family: 'IBM Plex Sans', sans-serif !important;
}}
textarea:focus {{
    border-color: {GUARD_GREEN} !important;
    box-shadow: 0 0 0 1px {GUARD_GREEN} !important;
}}
/* Tighten default vertical spacing between blocks so the Generate button
   is reachable without scrolling on most screens. Streamlit's vertical
   block container uses flexbox gap for spacing in recent versions.
   Kept deliberately simple/conservative after an earlier attempt with
   negative margins on individual elements caused overlap. */
div[data-testid="stVerticalBlock"] {{
    gap: 0.6rem !important;
}}
h1 {{
    font-family: 'Fraunces', serif !important;
    color: {INK} !important;
}}
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------------------------
# Header — native Streamlit components (title, badge, pain-point flag).
# Deliberately NOT raw HTML: some corporate networks/browsers strip inline
# HTML/CSS as an XSS precaution, which would leave this section blank.
# Native st.* calls always render regardless of any HTML sanitizer; the CSS
# block above only adds visual polish on top and degrades gracefully if lost.
# ---------------------------------------------------------------------------
st.markdown("# 🛡️ DoneGuard Lite")
st.caption("⚡ Powered by Claude AI + Groq")

st.warning(
    "**The pain point —** DoD checklists are usually generic, copy-pasted, "
    "and rarely tailored per Feature, so NFRs like security, accessibility, "
    "and performance quietly get skipped."
)

st.markdown(
    """
    **DoneGuard Lite** helps Agile Release Trains enforce **Built-In Quality**
    under SAFe. Paste a Feature's description and acceptance criteria below,
    and the app generates a tailored **Definition of Done checklist**,
    the **required role sign-offs**, and a **draft release note**.
    """
)
st.caption("💡 **How it works:** AI reads your Feature and generates all three, tailored to what you pasted.")

# ---------------------------------------------------------------------------
# API key — reads from Streamlit's Secrets so ANY visitor to this URL can
# use the app without entering their own key. You (the app owner) set this
# once in Streamlit Cloud's dashboard: App settings → Secrets, as:
#
#   GROQ_API_KEY = "your-real-key-here"
#
# Falls back to a sidebar input for local development/testing, or if no
# secret has been configured yet.
# ---------------------------------------------------------------------------
api_key = st.secrets.get("GROQ_API_KEY", None)

with st.sidebar:
    if api_key:
        st.markdown(f"""
        <div style="background-color:#EFEDE5; border-left:3px solid {GUARD_GREEN};
                    padding:10px 14px; border-radius:2px; font-size:0.85rem;
                    color:{INK};">
          ✅ Ready to go — no setup needed to try the app.
        </div>
        """, unsafe_allow_html=True)
    else:
        st.header("🔑 Groq API Key")
        api_key = st.text_input(
            "Enter your Groq API key",
            type="password",
            help="Your key is used only for this session and is never stored.",
        )
        st.caption(
            "Don't have a key? Get a free one (no credit card) at "
            "[console.groq.com/keys](https://console.groq.com/keys)."
        )

# ---------------------------------------------------------------------------
# Main screen — user input
# ---------------------------------------------------------------------------
st.subheader("Feature Details")

# Three varied sample Features so reviewers/judges can see the AI tailor
# its output differently depending on what kind of Feature it's given —
# not just repeat the same demo every time.
SAMPLE_FEATURES = {
    "🔐 Auth — Password Reset": """\
Feature: Allow customers to reset their password via email.

Acceptance Criteria:
1. User can request a password reset link from the login page by entering their registered email.
2. The reset link is sent via email and expires after 30 minutes.
3. Clicking an expired link shows a clear error and offers to resend.
4. New password must meet complexity rules (min 8 characters, 1 number, 1 symbol).
5. User receives a confirmation email once the password has been successfully changed.
6. All password reset attempts are logged for audit purposes.
""",
    "🔌 Backend — Payment API": """\
Feature: Expose a new REST API endpoint for processing customer refunds.

Acceptance Criteria:
1. Endpoint accepts an order ID and refund amount, validates against the original payment.
2. Partial refunds are supported, but total refunded cannot exceed the original charge.
3. Endpoint is authenticated via OAuth2 and rate-limited to prevent abuse.
4. A refund event is published to the downstream billing and analytics services.
5. Failed refunds return a clear error code and are retried up to 3 times automatically.
6. All refund transactions are logged with a correlation ID for support traceability.
""",
    "📊 Data — Sales Dashboard": """\
Feature: Add a real-time regional sales dashboard for store managers.

Acceptance Criteria:
1. Dashboard displays daily, weekly, and monthly sales totals filterable by region and store.
2. Data refreshes automatically every 5 minutes from the sales data warehouse.
3. Managers can only view data for stores within their assigned region (role-based access).
4. Dashboard must load within 3 seconds on a standard broadband connection.
5. Exportable as CSV for offline reporting.
6. Handles missing/delayed store data gracefully without breaking the dashboard.
""",
}

# Initialize the text area's session state once, so the sample-fill buttons
# can update it and have the change reflected in the widget below.
if "feature_text" not in st.session_state:
    st.session_state["feature_text"] = ""

def _load_sample(sample_text):
    st.session_state["feature_text"] = sample_text

def _clear_all():
    st.session_state["feature_text"] = ""

st.caption("Try a sample Feature:")
sample_cols = st.columns(len(SAMPLE_FEATURES) + 1)
for col, (label, text) in zip(sample_cols, SAMPLE_FEATURES.items()):
    col.button(label, on_click=_load_sample, args=(text,), use_container_width=True)
sample_cols[-1].button("🗑️ Clear", on_click=_clear_all, use_container_width=True)

feature_text = st.text_area(
    "📝 **Paste the Feature Description and Acceptance Criteria**",
    height=160,
    placeholder=(
        "e.g. Feature: Allow customers to reset their password via email.\n"
        "Acceptance Criteria:\n"
        "1. User can request a reset link from the login page.\n"
        "2. Reset link expires after 30 minutes.\n"
        "3. Password must meet complexity rules.\n"
        "..."
    ),
    key="feature_text",
)

generate_clicked = st.button("Generate Definition of Done", type="primary")

# Simple per-browser-session usage cap. Since this app now runs on a single
# shared API key (see above), this protects Groq's free-tier rate limits
# from being exhausted by one person spamming the button, while still
# giving each peer plenty of room to try the tool out.
MAX_GENERATIONS_PER_SESSION = 8
if "generation_count" not in st.session_state:
    st.session_state["generation_count"] = 0

# ---------------------------------------------------------------------------
# System prompt — instructs the model on exactly what to produce
# ---------------------------------------------------------------------------
SYSTEM_PROMPT = """\
You are an expert SAFe (Scaled Agile Framework) Release Train Engineer and
Quality Coach. Analyze the Feature's description and acceptance criteria,
then produce a CONCISE, SCANNABLE response — a busy RTE or judge should be
able to read the whole thing in under a minute. Prefer short phrases over
full sentences in checklist/list items. Do not restate the Feature text
back to the user.

Respond with EXACTLY three sections, using these Markdown headings verbatim:

## 1. Dynamic DoD Checklist
Markdown checkboxes ("- [ ] item"), max 8 items total, covering BOTH:
  - Technical / functional requirements specific to this Feature (max 4
    items — e.g. code complete, tests passing, code reviewed).
  - Non-Functional Requirements (NFRs) genuinely relevant to THIS Feature
    (max 4 items — choose only from security, accessibility, performance,
    data privacy, observability, scalability; skip any that don't apply).
  Keep each item to a single short line — no sub-explanations.

## 2. Required Sign-offs
Max 4 bullet points. Format each as "**Role** — one short reason (≤10 words)".
Only include roles that are genuinely relevant to this specific Feature.

## 3. Draft Release Notes
2-3 sentences max. Business-friendly, no jargon, no bullet points.

Be specific to the Feature given — never generic or boilerplate. If the
Feature text is vague, make reasonable assumptions and note them in ONE
short line at the end under "### Assumptions" (omit this section entirely
if no assumptions were needed).
"""


# ---------------------------------------------------------------------------
# Word document builder — converts the AI's Markdown-style output into a
# formatted .docx in memory (no disk writes), so it can be offered as a
# one-click download without any extra hosting/storage setup.
# ---------------------------------------------------------------------------
def _add_markdown_bold_runs(paragraph, text):
    """Split text on **bold** markers and add runs with matching formatting."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def build_docx(feature_text: str, result_text: str) -> bytes:
    doc = Document()

    title = doc.add_heading("DoneGuard Lite — Definition of Done", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x16, 0x23, 0x3D)

    doc.add_heading("Feature", level=2)
    doc.add_paragraph(feature_text.strip())

    for raw_line in result_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].lstrip("0123456789. "), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- [ ]") or line.startswith("- [x]"):
            p = doc.add_paragraph(style="List Bullet")
            checkbox = "☐ " if "[ ]" in line[:6] else "☑ "
            content = line.split("]", 1)[1].strip()
            p.add_run(checkbox)
            _add_markdown_bold_runs(p, content)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_markdown_bold_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_markdown_bold_runs(p, line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Logic — validate input and call the Groq API
# ---------------------------------------------------------------------------
if generate_clicked:
    if not api_key:
        st.warning("⚠️ Please enter your Groq API key in the sidebar first.")
    elif not feature_text.strip():
        st.warning("⚠️ Please paste a Feature description before generating.")
    elif st.session_state["generation_count"] >= MAX_GENERATIONS_PER_SESSION:
        st.warning(
            f"⚠️ You've reached the {MAX_GENERATIONS_PER_SESSION}-generation limit "
            "for this session (this keeps the shared demo key available for "
            "everyone). Refresh the page to reset, or use your own free Groq "
            "key in the sidebar for unlimited use."
        )
    else:
        try:
            client = Groq(api_key=api_key)

            with st.spinner("Analyzing Feature and generating compliance requirements..."):
                response = client.chat.completions.create(
                    model="openai/gpt-oss-120b",  # Groq's free-tier flagship model; swap for a lighter/faster model if you hit rate limits
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": feature_text},
                    ],
                    temperature=0.4,
                )

            result_text = response.choices[0].message.content
            st.session_state["generation_count"] += 1

            st.success("Definition of Done generated successfully!")
            st.markdown("---")
            st.markdown(result_text)

            st.download_button(
                "📥 Download as Word (.docx)",
                data=build_docx(feature_text, result_text),
                file_name="doneguard_definition_of_done.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            # Catches invalid API keys, network issues, rate limits, etc.
            st.error(f"❌ Something went wrong while calling the Groq API:\n\n{e}")
